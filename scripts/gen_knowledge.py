"""生成知识库的 docx/pdf 测试文档，验证格式兼容性。
运行：uv run --with python-docx --with reportlab scripts/gen_knowledge.py
"""
from pathlib import Path
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parent.parent
KNOWLEDGE = ROOT / "knowledge"

# 中文字体（Windows 自带宋体，reportlab 需要 ttf）
FONT_PATH = r"C:\Windows\Fonts\simsun.ttc"
if Path(FONT_PATH).exists():
    pdfmetrics.registerFont(TTFont("SimSun", FONT_PATH))


def gen_docx():
    doc = DocxDocument()
    doc.add_heading("出差管理制度", 0)
    doc.add_heading("一、出差审批", level=1)
    doc.add_paragraph("员工出差须提前 3 个工作日填写《出差申请单》，经部门负责人和分管领导审批后生效。")
    doc.add_heading("二、住宿标准", level=1)
    doc.add_paragraph("一线城市（北上广深）住宿标准为 500 元/晚，其他城市 350 元/晚，超出部分自理。")
    doc.add_heading("三、出差补贴", level=1)
    doc.add_paragraph("出差期间每日补贴 80 元，按实际出差天数计算，与当月工资一并发放。")
    doc.add_heading("四、差旅报销", level=1)
    doc.add_paragraph("出差结束 5 个工作日内提交报销，需附发票原件、行程单与出差申请单。")
    doc.add_heading("五、陪同客户", level=1)
    doc.add_paragraph("出差期间陪同客户产生的合理费用，凭发票与事由说明可另行报销。")
    out = KNOWLEDGE / "出差管理制度.docx"
    doc.save(str(out))
    print(f"生成 {out}")


def gen_pdf():
    out = KNOWLEDGE / "办公设备申领.pdf"
    c = canvas.Canvas(str(out), pagesize=A4)
    width, height = A4
    lines = [
        ("办公设备申领规范", 20),
        ("一、申领范围：笔记本电脑、显示器、键盘鼠标、耳机等办公设备。", 14),
        ("二、申领流程：新员工入职后由直属主管在 IT 服务平台提交设备申领单，", 14),
        ("IT 部门在 2 个工作日内完成设备发放。", 14),
        ("三、设备归还：员工离职前须将全部办公设备归还 IT 部门，", 14),
        ("损坏或丢失按设备残值赔偿。", 14),
        ("四、更换标准：笔记本使用满 3 年或出现影响使用的故障时可申请更换。", 14),
    ]
    y = height - 60
    for item in lines:
        if isinstance(item, tuple):
            text, size = item
        else:
            text, size = item, 14
        c.setFont("SimSun" if Path(FONT_PATH).exists() else "Helvetica", size)
        c.drawString(60, y, text)
        y -= 26
    c.save()
    print(f"生成 {out}")


if __name__ == "__main__":
    gen_docx()
    gen_pdf()
